"""
AI Document Detector — Text Extractor Service
Extracts text content and metadata from PDF and DOCX documents.
"""

from typing import Optional, Dict, Any, List
import io
import hashlib
from datetime import datetime


class TextExtractorService:
    """Service for extracting text and metadata from PDF and DOCX files."""

    def extract(self, file_bytes: bytes, file_ext: str) -> Optional[str]:
        """
        Extract text from a document file.

        Args:
            file_bytes: Raw file content as bytes.
            file_ext: File extension (e.g., '.pdf', '.docx').

        Returns:
            Extracted text string, or None if extraction failed.
        """
        if file_ext == ".pdf":
            return self._extract_from_pdf(file_bytes)
        elif file_ext == ".docx":
            return self._extract_from_docx(file_bytes)
        elif file_ext in (".png", ".jpg", ".jpeg"):
            # Images are handled by OCR service
            return None
        else:
            return None

    def _extract_from_pdf(self, file_bytes: bytes) -> Optional[str]:
        """Extract text from a PDF file using pdfplumber."""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            return "\n\n".join(text_parts) if text_parts else None

        except Exception as e:
            print(f"[TextExtractor] PDF extraction error: {e}")
            return None

    def _extract_from_docx(self, file_bytes: bytes) -> Optional[str]:
        """Extract text from a DOCX file using python-docx."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs) if paragraphs else None

        except Exception as e:
            print(f"[TextExtractor] DOCX extraction error: {e}")
            return None

    def extract_metadata(self, file_bytes: bytes, file_ext: str) -> Optional[Dict[str, Any]]:
        """
        Extract metadata from a document file for forensic analysis.

        Args:
            file_bytes: Raw file content as bytes.
            file_ext: File extension.

        Returns:
            Dictionary with metadata fields.
        """
        metadata: Dict[str, Any] = {
            "file_hash": hashlib.sha256(file_bytes).hexdigest(),
            "file_size_bytes": len(file_bytes),
        }

        if file_ext == ".pdf":
            pdf_meta = self._extract_pdf_metadata(file_bytes)
            if pdf_meta:
                metadata.update(pdf_meta)
        elif file_ext == ".docx":
            docx_meta = self._extract_docx_metadata(file_bytes)
            if docx_meta:
                metadata.update(docx_meta)

        return metadata

    def _extract_pdf_metadata(self, file_bytes: bytes) -> Optional[Dict[str, Any]]:
        """Extract metadata from PDF using pdfplumber."""
        try:
            import pdfplumber

            meta: Dict[str, Any] = {}
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                info = pdf.metadata or {}

                meta["author"] = info.get("Author", None)
                meta["creator_tool"] = info.get("Creator", None)
                meta["producer"] = info.get("Producer", None)
                meta["subject"] = info.get("Subject", None)
                meta["title"] = info.get("Title", None)
                meta["page_count"] = len(pdf.pages)

                # Parse dates
                for date_key in ["CreationDate", "ModDate"]:
                    raw_date = info.get(date_key, None)
                    if raw_date:
                        parsed = self._parse_pdf_date(raw_date)
                        if date_key == "CreationDate":
                            meta["creation_date"] = parsed
                        else:
                            meta["modification_date"] = parsed

                # Extract total word count
                total_words = 0
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        total_words += len(text.split())
                meta["total_words"] = total_words

                # Store raw metadata for forensics
                meta["raw"] = {k: str(v) for k, v in info.items()}

            return meta

        except Exception as e:
            print(f"[TextExtractor] PDF metadata error: {e}")
            return None

    def _extract_docx_metadata(self, file_bytes: bytes) -> Optional[Dict[str, Any]]:
        """Extract metadata from DOCX using python-docx."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            props = doc.core_properties

            meta: Dict[str, Any] = {
                "author": props.author,
                "creator_tool": props.last_modified_by,
                "title": props.title,
                "subject": props.subject,
                "category": props.category,
                "keywords": props.keywords,
                "revision": props.revision,
                "creation_date": props.created.isoformat() if props.created else None,
                "modification_date": props.modified.isoformat() if props.modified else None,
                "last_printed": props.last_printed.isoformat() if props.last_printed else None,
            }

            # Word and paragraph counts
            total_words = 0
            total_paragraphs = 0
            for p in doc.paragraphs:
                if p.text.strip():
                    total_paragraphs += 1
                    total_words += len(p.text.split())

            meta["total_words"] = total_words
            meta["total_paragraphs"] = total_paragraphs
            meta["page_count"] = None  # DOCX doesn't have reliable page count

            # Store raw
            meta["raw"] = {
                "author": props.author,
                "last_modified_by": props.last_modified_by,
                "revision": str(props.revision) if props.revision else None,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
            }

            return meta

        except Exception as e:
            print(f"[TextExtractor] DOCX metadata error: {e}")
            return None

    def _parse_pdf_date(self, date_str: str) -> Optional[str]:
        """
        Parse PDF date string format: D:YYYYMMDDHHmmSS+HH'mm'
        Returns ISO format string.
        """
        try:
            if not date_str:
                return None

            # Remove 'D:' prefix
            clean = date_str.replace("D:", "").strip()

            # Try to parse common PDF date formats
            for fmt in [
                "%Y%m%d%H%M%S",
                "%Y%m%d%H%M",
                "%Y%m%d",
                "%Y",
            ]:
                try:
                    # Remove timezone info for parsing
                    date_part = clean[:len(fmt.replace("%", "").replace("Y", "0").replace("m", "0").replace("d", "0").replace("H", "0").replace("M", "0").replace("S", "0"))]
                    dt = datetime.strptime(clean[:14].ljust(14, "0"), "%Y%m%d%H%M%S")
                    return dt.isoformat()
                except (ValueError, IndexError):
                    continue

            return date_str  # Return raw if parsing fails

        except Exception:
            return date_str
